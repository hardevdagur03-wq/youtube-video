"""DOCX Exporter — Phase 10.

Generates Microsoft Word documents with proper styles, TOC, images, tables,
code blocks, headers/footers, and professional formatting.
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from models.blog_export import ExportRequest, ExportFile
from export.base import BaseExporter


class DOCXExporter(BaseExporter):
    """Exports blog to Microsoft Word .docx format."""

    def format_name(self) -> str:
        return "docx"

    def file_extension(self) -> str:
        return ".docx"

    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def export(self, request: ExportRequest, output_dir: Path) -> ExportFile:
        filename = f"{self.sanitize_filename(request.blog_title)}{self.file_extension()}"
        filepath = output_dir / filename

        doc = Document()
        self._setup_styles(doc)

        # Cover page metadata
        if request.blog_title:
            title = doc.add_heading(request.blog_title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if request.author:
            meta.add_run(f"By {request.author}").bold = True
            meta.add_run("\n")
        if request.publish_date:
            meta.add_run(request.publish_date)
            meta.add_run("\n")
        if request.reading_time:
            meta.add_run(f"{request.reading_time} · {request.word_count:,} words")
        doc.add_page_break()

        # Table of Contents
        if request.table_of_contents:
            doc.add_heading("Table of Contents", level=1)
            for item in request.table_of_contents:
                p = doc.add_paragraph(item, style='List Bullet')
            doc.add_page_break()

        # Introduction
        if request.introduction:
            self._add_html_text(doc, request.introduction, is_intro=True)

        # Sections
        for section in request.sections:
            heading = section.get("heading", "")
            content = section.get("content", "")
            subsections = section.get("subsections", [])
            callouts = section.get("callout_boxes", [])

            if heading:
                doc.add_heading(heading, level=1)
            if content:
                self._add_html_text(doc, content)

            for callout in callouts:
                ct = callout.get("type", "note")
                ct_text = callout.get("text", "")
                p = doc.add_paragraph()
                run = p.add_run(f"[{ct.upper()}] {ct_text}")
                run.italic = True
                run.font.size = Pt(10)

            for sub in subsections:
                sub_heading = sub.get("heading", "")
                sub_content = sub.get("content", "")
                if sub_heading:
                    doc.add_heading(sub_heading, level=2)
                if sub_content:
                    self._add_html_text(doc, sub_content)

        # FAQ
        if request.faq:
            doc.add_page_break()
            doc.add_heading("Frequently Asked Questions", level=1)
            for faq in request.faq:
                q = faq.get("question", "")
                a = faq.get("answer", "")
                if q:
                    p = doc.add_paragraph()
                    run = p.add_run(q)
                    run.bold = True
                if a:
                    doc.add_paragraph(a)

        # Conclusion
        if request.conclusion:
            doc.add_heading("Conclusion", level=1)
            self._add_html_text(doc, request.conclusion)

        # CTA
        if request.call_to_action:
            p = doc.add_paragraph()
            run = p.add_run(request.call_to_action)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)

        # References
        if request.references:
            doc.add_heading("References", level=1)
            for ref in request.references:
                label = ref.get("label", ref.get("url", ""))
                url = ref.get("url", "")
                p = doc.add_paragraph()
                if label and url:
                    p.add_run(f"{label}: ").bold = True
                    self._add_hyperlink(p, url, url)
                elif url:
                    self._add_hyperlink(p, url, url)
                elif label:
                    p.add_run(label)

        # Footer with page numbers
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"{request.blog_title} — Page ").font.size = Pt(8)
            # Page number field
            run = p.add_run()
            fld_char1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
            run._r.append(fld_char1)
            run2 = p.add_run()
            instr = run2._r.makeelement(qn('w:instrText'), {})
            instr.text = ' PAGE '
            run2._r.append(instr)
            run3 = p.add_run()
            fld_char2 = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
            run3._r.append(fld_char2)

        doc.save(str(filepath))
        size = filepath.stat().st_size
        return ExportFile(
            format=self.format_name(),
            filename=filename,
            size_bytes=size,
            size_display=self._size_display(size),
            download_url=f"/api/export/download/{filename}",
            mime_type=self.mime_type(),
        )

    def _setup_styles(self, doc: Document) -> None:
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        for level in range(1, 4):
            hs = doc.styles[f'Heading {level}']
            hs.font.name = 'Calibri'
            hs.font.color.rgb = RGBColor(0x1F, 0x29, 0x3D)
            if level == 1:
                hs.font.size = Pt(18)
            elif level == 2:
                hs.font.size = Pt(14)
            else:
                hs.font.size = Pt(12)

    def _add_html_text(self, doc: Document, text: str, is_intro: bool = False) -> None:
        """Add text with basic formatting (bold, italic, code, links, lists)."""
        import re
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Unordered list
            if para_text.startswith('- ') or para_text.startswith('* '):
                for line in para_text.split('\n'):
                    line = line.strip()
                    if line.startswith('- ') or line.startswith('* '):
                        content = line[2:]
                        doc.add_paragraph(content, style='List Bullet')
                continue

            # Ordered list
            if re.match(r'^\d+\.\s', para_text):
                for line in para_text.split('\n'):
                    line = line.strip()
                    m = re.match(r'^\d+\.\s(.+)$', line)
                    if m:
                        doc.add_paragraph(m.group(1), style='List Number')
                continue

            # Code block
            if para_text.startswith('```'):
                code_lines = []
                for line in para_text.split('\n'):
                    if line.startswith('```'):
                        continue
                    code_lines.append(line)
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(1)
                continue

            # Regular paragraph with inline formatting
            p = doc.add_paragraph()
            self._add_inline_formatting(p, para_text)

    def _add_inline_formatting(self, paragraph, text: str) -> None:
        import re
        # Process inline code, bold, italic, links
        parts = re.split(r'(```[^`]*```|`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\))', text)
        for part in parts:
            if not part:
                continue
            # Inline code
            if part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part.strip('`'))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
            # Bold
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part.strip('*'))
                run.bold = True
            # Italic
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = paragraph.add_run(part.strip('*'))
                run.italic = True
            # Link
            elif part.startswith('[') and '](' in part:
                m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
                if m:
                    self._add_hyperlink(paragraph, m.group(2), m.group(1))
            else:
                paragraph.add_run(part)

    def _add_hyperlink(self, paragraph, url: str, text: str) -> None:
        part = paragraph.add_run(text)
        part.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)
        part.underline = True
        # Add hyperlink relationship
        try:
            parent = paragraph._p
            hyperlink = parent.makeelement(qn('w:hyperlink'), {
                qn('r:id'): '', qn('w:history'): '1'
            })
            # We can't easily add real hyperlinks without document relationships
            # so we use colored underlined text as a visual indicator
        except Exception:
            pass
